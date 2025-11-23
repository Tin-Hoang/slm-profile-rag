"""Document processing module for loading and chunking documents."""

import logging
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .config_loader import get_config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats and chunk them for RAG."""

    def __init__(self):
        """Initialize document processor with configuration."""
        self.config = get_config()
        self.chunk_size = self.config.get("document_processing.chunk_size", 1000)
        self.chunk_overlap = self.config.get("document_processing.chunk_overlap", 200)
        self.supported_extensions = self.config.get(
            "document_processing.supported_extensions",
            [".pdf", ".docx", ".doc", ".html", ".htm", ".txt", ".md"],
        )

        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def load_pdf(self, file_path: Path) -> list[Document]:
        """Load and extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            List of Document objects
        """
        try:
            loader = PyPDFLoader(str(file_path))
            documents = loader.load()
            logger.info(f"Loaded {len(documents)} pages from {file_path.name}")
            return documents
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return []

    def load_docx(self, file_path: Path) -> list[Document]:
        """Load and extract text from Word document.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of Document objects
        """
        try:
            doc = DocxDocument(str(file_path))
            text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text])

            # Extract tables if present
            if doc.tables:
                for table in doc.tables:
                    table_text = "\n".join(
                        ["\t".join([cell.text for cell in row.cells]) for row in table.rows]
                    )
                    text += f"\n\n{table_text}"

            document = Document(
                page_content=text,
                metadata={"source": str(file_path), "type": "docx"},
            )
            logger.info(f"Loaded Word document {file_path.name}")
            return [document]
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return []

    def load_html(self, file_path: Path) -> list[Document]:
        """Load and extract text from HTML file.

        Args:
            file_path: Path to HTML file

        Returns:
            List of Document objects
        """
        try:
            with open(file_path, encoding="utf-8") as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, "lxml")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text(separator="\n")

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            text = "\n".join(line for line in lines if line)

            document = Document(
                page_content=text,
                metadata={"source": str(file_path), "type": "html"},
            )
            logger.info(f"Loaded HTML file {file_path.name}")
            return [document]
        except Exception as e:
            logger.error(f"Error loading HTML {file_path}: {e}")
            return []

    def load_text(self, file_path: Path) -> list[Document]:
        """Load text file (txt, md, etc).

        Args:
            file_path: Path to text file

        Returns:
            List of Document objects
        """
        try:
            loader = TextLoader(str(file_path), encoding="utf-8")
            documents = loader.load()
            logger.info(f"Loaded text file {file_path.name}")
            return documents
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return []

    def load_document(self, file_path: Path) -> list[Document]:
        """Load document based on file extension.

        Args:
            file_path: Path to document

        Returns:
            List of Document objects
        """
        extension = file_path.suffix.lower()

        if extension == ".pdf":
            return self.load_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self.load_docx(file_path)
        elif extension in [".html", ".htm"]:
            return self.load_html(file_path)
        elif extension in [".txt", ".md"]:
            return self.load_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {extension} for {file_path.name}")
            return []

    def chunk_documents(self, documents: list[Document]) -> list[Document]:
        """Split documents into chunks.

        Args:
            documents: List of Document objects

        Returns:
            List of chunked Document objects
        """
        if not documents:
            return []

        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Split {len(documents)} documents into {len(chunks)} chunks")
            return chunks
        except Exception as e:
            logger.error(f"Error chunking documents: {e}")
            return []

    def process_directory(self, directory_path: str | Path) -> list[Document]:
        """Process all documents in a directory.

        Args:
            directory_path: Path to directory containing documents

        Returns:
            List of chunked Document objects
        """
        directory = Path(directory_path)

        if not directory.exists():
            logger.error(f"Directory not found: {directory}")
            return []

        # Get main document path to exclude it from vector store
        main_doc_path = self.config.get("main_document.path", "")
        main_doc_file = Path(main_doc_path).resolve() if main_doc_path else None

        all_documents = []
        skipped_main_doc = False

        # Find all supported files
        for ext in self.supported_extensions:
            files = list(directory.rglob(f"*{ext}"))
            for file_path in files:
                # Skip main document file - it's loaded directly, not chunked
                if main_doc_file and file_path.resolve() == main_doc_file:
                    if not skipped_main_doc:
                        logger.info(
                            f"Skipping main document: {file_path.name} "
                            "(loaded directly, not stored in vector DB)"
                        )
                        skipped_main_doc = True
                    continue

                logger.info(f"Processing {file_path.name}...")
                docs = self.load_document(file_path)
                all_documents.extend(docs)

        if not all_documents:
            logger.warning(f"No documents found in {directory}")
            return []

        logger.info(f"Loaded {len(all_documents)} documents from {directory}")

        # Chunk all documents
        chunked_documents = self.chunk_documents(all_documents)

        return chunked_documents


def process_documents(directory_path: str | None = None) -> list[Document]:
    """Convenience function to process documents.

    Args:
        directory_path: Path to documents directory (uses config default if None)

    Returns:
        List of chunked Document objects
    """
    config = get_config()
    if directory_path is None:
        directory_path = config.get_env("DOCUMENTS_DIR", "./data/documents")

    processor = DocumentProcessor()
    return processor.process_directory(directory_path)
